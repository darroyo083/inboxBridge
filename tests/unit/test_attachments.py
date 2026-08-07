"""Unit tests for gmail/attachments.py — fixtures generated in-memory."""

from __future__ import annotations

import io
from typing import Any

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter

from inboxbridge.config import Settings
from inboxbridge.gmail.attachments import (
    AttachmentPasswordError,
    extract_attachment_text,
    extract_attachments,
)
from inboxbridge.gmail.parse import RawAttachment
from inboxbridge.models import AttachmentMeta

PDF_PASSWORD = "s3cret"


def make_settings(**overrides: object) -> Settings:
    base: dict[str, Any] = {
        "_env_file": None,
        "attachment_max_bytes": 10 * 1024 * 1024,
        "attachment_max_text_chars": 20_000,
        "attachment_max_count": 5,
        "PDF_PASSWORD": PDF_PASSWORD,
    }
    base.update(overrides)
    return Settings(**base)


def make_pdf_bytes(text: str) -> bytes:
    """Hand-rolled minimal PDF with one text object (no external tooling)."""
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 11 Tf 50 750 Td ({escaped}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    out += b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:])
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    ).encode()
    return bytes(out)


def make_encrypted_pdf_bytes(text: str, password: str) -> bytes:
    reader = PdfReader(io.BytesIO(make_pdf_bytes(text)))
    writer = PdfWriter(clone_from=reader)
    writer.encrypt(password)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_docx_bytes(paragraphs: list[str], rows: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    for row in rows or []:
        table = doc.add_table(rows=1, cols=len(row))
        for index, value in enumerate(row):
            table.rows[0].cells[index].text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def att(filename: str, mime: str, content: bytes) -> RawAttachment:
    return RawAttachment(
        filename=filename, mime_type=mime, size_bytes=len(content), content=content
    )


class TestPDF:
    def test_extracts_text_layer(self) -> None:
        meta = extract_attachment_text(
            att("invoice.pdf", "application/pdf", make_pdf_bytes("Invoice 2024 total 42.50 EUR")),
            make_settings(),
        )
        assert meta.extracted_text == "Invoice 2024 total 42.50 EUR"
        assert meta.filename == "invoice.pdf"

    def test_wrong_password_raises_typed_error(self) -> None:
        content = make_encrypted_pdf_bytes("Secret content", "real-pw")
        with pytest.raises(AttachmentPasswordError):
            extract_attachment_text(att("locked.pdf", "application/pdf", content), make_settings())

    def test_correct_password_extracts(self) -> None:
        content = make_encrypted_pdf_bytes("Secret content", PDF_PASSWORD)
        meta = extract_attachment_text(
            att("locked.pdf", "application/pdf", content), make_settings()
        )
        assert meta.extracted_text == "Secret content"

    def test_missing_password_raises(self) -> None:
        content = make_encrypted_pdf_bytes("Secret content", "real-pw")
        settings = make_settings(PDF_PASSWORD="")
        with pytest.raises(AttachmentPasswordError):
            extract_attachment_text(att("locked.pdf", "application/pdf", content), settings)


class TestDocx:
    def test_extracts_paragraphs_and_tables(self) -> None:
        content = make_docx_bytes(
            ["Hola, este es el informe.", "Segunda línea."], rows=[["A", "B", "C"]]
        )
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        meta = extract_attachment_text(att("report.docx", mime, content), make_settings())
        assert "Hola, este es el informe." in meta.extracted_text
        assert "A | B | C" in meta.extracted_text

    def test_non_docx_bytes_yield_empty_text(self) -> None:
        meta = extract_attachment_text(
            att("fake.docx", "application/zip", b"PK\x03\x04 not a docx"), make_settings()
        )
        assert meta.extracted_text == ""

    def test_doc_binary_format_yields_empty_text(self) -> None:
        meta = extract_attachment_text(
            att("old.doc", "application/msword", b"\xd0\xcf\x11\xe0 old doc"), make_settings()
        )
        assert meta.extracted_text == ""


class TestTxt:
    def test_utf8(self) -> None:
        meta = extract_attachment_text(
            att("notes.txt", "text/plain", "Grüße – Factura 2025".encode()), make_settings()
        )
        assert meta.extracted_text == "Grüße – Factura 2025"

    def test_latin1_fallback(self) -> None:
        content = "café résumé".encode("latin-1")
        meta = extract_attachment_text(att("notes.txt", "text/plain", content), make_settings())
        assert meta.extracted_text == "café résumé"

    def test_mime_only_txt(self) -> None:
        meta = extract_attachment_text(
            att("data", "text/plain", b"plain content"), make_settings()
        )
        assert meta.extracted_text == "plain content"


class TestLimits:
    def test_too_large_skips_text_but_keeps_meta(self) -> None:
        settings = make_settings(attachment_max_bytes=10)
        meta = extract_attachment_text(
            att("big.pdf", "application/pdf", make_pdf_bytes("too big")), settings
        )
        assert meta.extracted_text == ""
        assert meta.filename == "big.pdf"
        assert meta.size_bytes > 10

    def test_text_truncated_to_char_limit(self) -> None:
        settings = make_settings(attachment_max_text_chars=10)
        meta = extract_attachment_text(
            att("long.txt", "text/plain", b"a" * 100), settings
        )
        assert meta.extracted_text == "a" * 10

    def test_count_limit(self) -> None:
        settings = make_settings(attachment_max_count=2)
        items = [
            att(f"f{i}.txt", "text/plain", f"content {i}".encode()) for i in range(5)
        ]
        metas = extract_attachments(items, settings)
        assert [m.filename for m in metas] == ["f0.txt", "f1.txt"]

    def test_unsupported_type_yields_empty_text(self) -> None:
        meta = extract_attachment_text(
            att("book.xlsx", "application/vnd.ms-excel", b"binary"), make_settings()
        )
        assert meta.extracted_text == ""

    def test_wrong_password_propagates_from_batch(self) -> None:
        locked = att("locked.pdf", "application/pdf", make_encrypted_pdf_bytes("x", "real-pw"))
        with pytest.raises(AttachmentPasswordError):
            extract_attachments([locked], make_settings())

    def test_meta_type_is_frozen_attachment_meta(self) -> None:
        meta = extract_attachment_text(att("a.txt", "text/plain", b"hi"), make_settings())
        assert isinstance(meta, AttachmentMeta)
