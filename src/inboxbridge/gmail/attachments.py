"""Local attachment text extraction: PDF (text layer only), DOCX, TXT.

Content stays in memory and is NEVER persisted. Files are untrusted data:
parsed as text only, never executed (no macros, no embedded scripts run).
No OCR — scanned PDFs yield no text.

Limits come from config: ``attachment_max_bytes`` (per file), 
``attachment_max_text_chars`` (per file) and ``attachment_max_count``.
Oversized or unsupported files still yield an AttachmentMeta so the pipeline
knows the attachment existed, but with empty ``extracted_text``. A wrong PDF
password raises the typed :class:`AttachmentPasswordError` (the operator
must fix the password or resend the file).
"""

from __future__ import annotations

import io
import logging

from docx import Document
from pypdf import PdfReader
from pypdf.errors import WrongPasswordError

from ..config import Settings
from ..models import AttachmentMeta
from .parse import RawAttachment

logger = logging.getLogger(__name__)


class AttachmentError(RuntimeError):
    """Base class for attachment extraction failures."""


class AttachmentPasswordError(AttachmentError):
    """The PDF is password-protected and the configured password is wrong/absent."""


def extract_attachment_text(att: RawAttachment, settings: Settings) -> AttachmentMeta:
    """Extract text from one attachment, honoring size/char limits."""
    if att.size_bytes > settings.attachment_max_bytes:
        logger.info("attachment too large, no text: %s (%d bytes)", att.filename, att.size_bytes)
        return AttachmentMeta(
            filename=att.filename, mime_type=att.mime_type, size_bytes=att.size_bytes
        )
    text = _extract_text(att, settings)
    if len(text) > settings.attachment_max_text_chars:
        text = text[: settings.attachment_max_text_chars]
        logger.info(
            "attachment text truncated to %d chars: %s",
            settings.attachment_max_text_chars,
            att.filename,
        )
    return AttachmentMeta(
        filename=att.filename,
        mime_type=att.mime_type,
        size_bytes=att.size_bytes,
        extracted_text=text,
    )


def extract_attachments(atts: list[RawAttachment], settings: Settings) -> list[AttachmentMeta]:
    """Extract text for up to ``attachment_max_count`` attachments.

    Wrong PDF passwords raise :class:`AttachmentPasswordError` (typed, so the
    caller can notify the operator); everything else degrades to empty text.
    """
    limit = settings.attachment_max_count
    selected = atts[:limit]
    if len(atts) > limit:
        logger.warning(
            "attachment count %d exceeds limit %d; ignoring %d",
            len(atts),
            limit,
            len(atts) - limit,
        )
    return [extract_attachment_text(att, settings) for att in selected]


def _extract_text(att: RawAttachment, settings: Settings) -> str:
    try:
        if _is_pdf(att):
            return _pdf_text(att, settings)
        if _is_docx(att):
            return _docx_text(att)
        if _is_txt(att):
            return _txt_text(att)
        return ""
    except AttachmentPasswordError:
        raise
    except Exception:
        logger.exception("attachment extraction failed, no text: %s", att.filename)
        return ""


def _is_pdf(att: RawAttachment) -> bool:
    return att.mime_type == "application/pdf" or att.filename.lower().endswith(".pdf")


def _is_docx(att: RawAttachment) -> bool:
    return att.filename.lower().endswith(".docx")


def _is_txt(att: RawAttachment) -> bool:
    return att.filename.lower().endswith(".txt") or att.mime_type == "text/plain"


def _pdf_text(att: RawAttachment, settings: Settings) -> str:
    try:
        reader = PdfReader(io.BytesIO(att.content))
        if reader.is_encrypted:
            password = settings.pdf_password.get_secret_value()
            if not password:
                raise AttachmentPasswordError(
                    f"PDF {att.filename} is password-protected (set PDF_PASSWORD)"
                )
            if reader.decrypt(password) == 0:
                raise AttachmentPasswordError(
                    f"wrong PDF password for {att.filename} (set PDF_PASSWORD)"
                )
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except WrongPasswordError as exc:
        raise AttachmentPasswordError(f"wrong PDF password for {att.filename}") from exc


def _docx_text(att: RawAttachment) -> str:
    try:
        doc = Document(io.BytesIO(att.content))
    except Exception:
        return ""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _txt_text(att: RawAttachment) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return att.content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return att.content.decode("utf-8", errors="replace")
